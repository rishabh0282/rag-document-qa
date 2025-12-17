"""
Document loader supporting PDF, TXT, DOCX with metadata.
"""
from dataclasses import dataclass
from typing import List, Optional, Dict, Any
import logging
import os

from PyPDF2 import PdfReader
import docx

logger = logging.getLogger(__name__)


@dataclass
class Document:
    id: str
    filename: str
    text: str
    metadata: Dict[str, Any]


class DocumentLoader:
    def __init__(self) -> None:
        pass

    def _load_txt(self, path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            return fh.read()

    def _load_pdf(self, path: str) -> str:
        try:
            reader = PdfReader(path)
            pages = []
            for i, p in enumerate(reader.pages):
                try:
                    pages.append(p.extract_text() or "")
                except Exception:
                    logger.exception("Failed to extract page %s from %s", i, path)
                    pages.append("")
            return "\n".join(pages)
        except Exception:
            logger.exception("Failed to read PDF %s", path)
            raise

    def _load_docx(self, path: str) -> str:
        try:
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            logger.exception("Failed to read DOCX %s", path)
            raise

    def load(self, path: str, doc_id: Optional[str] = None) -> Document:
        if not os.path.exists(path):
            raise FileNotFoundError(path)
        ext = os.path.splitext(path)[1].lower()
        text = ""
        if ext == ".pdf":
            text = self._load_pdf(path)
        elif ext in (".txt", ".md"):
            text = self._load_txt(path)
        elif ext in (".docx",):
            text = self._load_docx(path)
        else:
            raise ValueError(f"Unsupported extension: {ext}")
        metadata = {"filename": os.path.basename(path)}
        return Document(id=doc_id or metadata["filename"], filename=metadata["filename"], text=text, metadata=metadata)

    def load_batch(self, paths: List[str]) -> List[Document]:
        docs = []
        for p in paths:
            try:
                docs.append(self.load(p))
            except Exception:
                logger.exception("Skipping file due to error: %s", p)
        return docs